from __future__ import unicode_literals

from lxml import etree
import docx
import cStringIO as StringIO
import zipfile
import os
import bs4

CURRENT_PATH = os.path.dirname(os.path.realpath(__file__ ))

# A hack for creating the base style and numbering files.
# Eventually replace this by generating the numbering and style files from scratch.
INITIAL_NUMID = 13
INITIAL_ABSTRACT_NUMID = 12

def generate_wordrelationships(relationshiplist):
	'''Generate a Word relationships file'''
	# Default list of relationships
	# FIXME: using string hack instead of making element
	#relationships = makeelement('Relationships', nsprefix='pr')
	relationships = etree.fromstring(
		'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006'
		'/relationships"></Relationships>')
	count = 0
	for n, relationship in relationshiplist:
		# Relationship IDs (rId) start at 1.
		a = {
			'Id':     'rId'+str(n),
			'Type':   relationship[0],
			'Target': relationship[1]
		}
		try:
			a['TargetMode'] = relationship[2]
		except IndexError:
			pass

		# print a

		rel_elm = docx.makeelement('Relationship', nsprefix=None,
							  attributes=a
							  )
		relationships.append(rel_elm)
		count += 1
	return relationships


class TrackableObject(object):
	_latest_id = 0

	def __init__(self, document=None):
		if not hasattr(type(self), "_cache"):
			type(self)._cache = {}

		type(self)._latest_id += 1
		self.id = type(self)._latest_id
		self._add_to_list(self)

		self.document = document

	def namespace_string(self, s):
		return self.document.namespace_string(s)

	def _add_to_list(self, instance):
		self._cache[instance.id] = instance


class AbstractNum(TrackableObject):
	# Add an abstractNum element to the numbering.xml file.
	# Describes the properties of a numberedList.

	_latest_id = 11

	BULLET_FORMAT = "bullet"
	DECIMAL_FORMAT = "decimal"
	LOWER_LETTER_FORMAT = "lowerLetter"
	UPPER_LETTER_FORMAT = "upperLetter"
	LOWER_ROMAN_FORMAT = "lowerRoman"
	UPPER_ROMAN_FORMAT = "upperRoman"

	def __init__(self, *args, **kwargs):
		self.levels = []
		return super(AbstractNum, self).__init__(*args, **kwargs)

	def add_level(self, level, start_from=1, number_format=UPPER_LETTER_FORMAT, text_format="%1.", justification="left", text_indentation=360, bullet_distance=360):
		# ilvl is the indentation level of the list (zero-indexed).
		level = {
			'ilvl': level,
			'_subelements': [],
		}
		level_subelements = level['_subelements']
		# start is the number from which the list will start.
		level_subelements.append(('start', {'val': start_from}))
		# This specifies how the number itself is displayed.
		level_subelements.append(('numFmt', {'val': number_format}))
		# This specifies how the bullet text is displayed.
		level_subelements.append(('lvlText', {'val': text_format}))
		# This specifies the level justification.
		level_subelements.append(('lvlJc', {'val': justification}))
		# The paragraph properties element is used to describe the indentation of the list.
		# For ind, left describes the distance between the text and the left margin of the page.
		# hanging describes the distance between the left edge of the text and the bullet.
		level_subelements.append(('pPr', {
			'_subelements': [
				('tabs', {
					'_subelements': [
						('tab', {'val': 'num', 'pos': text_indentation})
					]
				}),
				('ind', {'left': text_indentation, 'hanging': bullet_distance}),
			],
		}))

		self.levels.append(level)

	def as_element(self, parent):
		abstractNum_element = etree.SubElement(parent, self.namespace_string("abstractNum"))
		abstractNum_element.attrib[self.namespace_string("abstractNumId")] = str(self.id)
		self.levels.sort(key=lambda l: l['ilvl'])

		def create_element(parent_element, element_name, element_properties):
			e = etree.SubElement(parent_element, self.namespace_string(element_name))
			subelements = element_properties.pop('_subelements', [])
			for name, value in element_properties.items():
				e.attrib[self.namespace_string(name)] = str(value)

			for name, properties in subelements:
				create_element(e, name, properties)


		for level in self.levels:
			create_element(abstractNum_element, 'lvl', level)

		return abstractNum_element


class Num(TrackableObject):
	_latest_id = 13

	# Add a num element to the numbering.xml file.
	# Describes a specific instance of a number based on an abstractNum element.

	def __init__(self, *args, **kwargs):
		abstractNum = kwargs.pop('abstractNum')
		self.abstractNumId = abstractNum.id
		super(Num, self).__init__(*args, **kwargs)

	def as_element(self, parent):
		num_element = etree.SubElement(parent, self.namespace_string("num"))
		num_element.attrib[self.namespace_string("numId")] = str(self.id)
		# A reference to the abstractNum element which specifies the properties of this particular num.
		abstractNumId_element = etree.SubElement(num_element, self.namespace_string("abstractNumId"))
		abstractNumId_element.attrib[self.namespace_string("val")] = str(self.abstractNumId)
		# A lvlOverride element overides some of the properties of the lvl for that abstractNum.
		lvlOverride_element = etree.SubElement(num_element, self.namespace_string("lvlOverride"))
		lvlOverride_element.attrib[self.namespace_string("ilvl")] = "0"
		# We override the starting number to make sure the list always starts from 1.
		startOverride_element = etree.SubElement(lvlOverride_element, self.namespace_string("startOverride"))
		startOverride_element.attrib[self.namespace_string("val")] = "1"


class Style(object):
	PARAGRAPH_TYPE = "paragraph"
	DEFAULT_TYPE = PARAGRAPH_TYPE

	def __init__(self, document=None, name="", *args, **kwargs):
		self.document = document
		self.type = self.DEFAULT_TYPE
		self.name = name
		self.numId = None
		super(Style, self).__init__(*args, **kwargs)

	def namespace_string(self, s):
		return self.document.namespace_string(s)

	def set_number_id(self, numId):
		self.numId = numId

	def as_element(self, parent):
		style_element = etree.SubElement(parent, self.namespace_string("style"))
		style_element.attrib[self.namespace_string("styleId")] = self.name
		style_element.attrib[self.namespace_string("type")] = self.type
		if self.numId:
			pPr_element = etree.SubElement(style_element, self.namespace_string("pPr"))
			numPr_element = etree.SubElement(pPr_element, self.namespace_string("numPr"))
			numId_element = etree.SubElement(numPr_element, self.namespace_string("numId"))
			numId_element.attrib[self.namespace_string('val')] = str(self.numId)
		etree.SubElement(pPr_element, self.namespace_string("contextualSpacing"))


class DocumentElement(object):
	def __init__(self, document=None):
		self.document = document

	def namespace_string(self, tag):
		return self.document.namespace_string(tag)

	def add_styles(self, style_tree, numbering_tree, num=None, abstractNum=None):
		pass

	def get_elements(self):
		pass


class Heading(DocumentElement):
	DEFAULT_LEVEL = 1

	def __init__(self, text, level=DEFAULT_LEVEL, *args, **kwargs):
		self.text = text
		self.level = level

		super(Heading,self).__init__(*args, **kwargs)

	def get_elements(self):
		return [docx.heading(self.text, self.level), ]


class RunProperties(DocumentElement):
	def __init__(self, properties):
		self.properties = properties

	def as_element(self):
		run_element = docx.makeelement('rPr')
		for k, v in self.properties.items():
			if v:
				property_element = docx.makeelement(k, attributes={'val': v})
			else:
				property_element = docx.makeelement(k)
			run_element.append(property_element)

		return run_element


class Paragraph(DocumentElement):
	DEFAULT_STYLE = "BodyText"

	def __init__(self, text=None, style=DEFAULT_STYLE, num_level=None, num_id=None, *args, **kwargs):
		# self.text = text
		self.style = style
		self.hyperlink_text = ""
		self.hyperlink_uri = ""
		self.num_level = num_level
		self.num_id = num_id

		self.elements = []

		if text is not None: self.elements.append((text, None))

		super(Paragraph, self).__init__(*args, **kwargs)

	def add_text(self, text, italic=False, bold=False, superscript=False, subscript=False):
		props = {}
		if italic: props['i'] = None
		if bold: props['b'] = None
		if superscript: props['vertAlign'] = 'superscript'
		if subscript: props['vertAlign'] = 'subscript'
		if props:
			self.elements.append((text, RunProperties(props)))
		else:
			self.elements.append((text, None))

	def get_elements(self):
		if not self.elements:
			paragraph_element = docx.paragraph(self.text, self.style)
			if self.hyperlink_text:
				r_element = paragraph_element.find(self.namespace_string("r"))
				t_element = r_element.find(self.namespace_string("t"))
				t_element.attrib["{http://www.w3.org/XML/1998/namespace}space"] = "preserve"
				hyperlink_element = docx.makeelement('hyperlink', attributes={'id': 'rId%s' % self.hyperlinkId}, attrnsprefix='r')
				rPr_element = docx.makeelement('rPr')
				rStyle_element = docx.makeelement('rStyle', attributes={'val': 'Hyperlink'})
				r_element = docx.makeelement('r')
				t_element = docx.makeelement('t', tagtext=self.hyperlink_text)
				rPr_element.append(rStyle_element)
				r_element.append(rPr_element)
				r_element.append(t_element)
				hyperlink_element.append(r_element)

				paragraph_element.append(hyperlink_element)
		else:
			paragraph_element = docx.makeelement('p')
			pPr_element = docx.makeelement('pPr')
			pStyle_element = docx.makeelement('pStyle', attributes={'val': self.style})
			pPr_element.append(pStyle_element)
			if self.num_level is not None and self.num_id is not None:
				numPr_element = docx.makeelement('numPr')
				ilvl_element = docx.makeelement("ilvl", attributes={'val': str(self.num_level)})
				numId_element = docx.makeelement("numId", attributes={'val': str(self.num_id)})
				numPr_element.append(ilvl_element)
				numPr_element.append(numId_element)
				pPr_element.append(numPr_element)
			paragraph_element.append(pPr_element)

			for text, properties in self.elements:
				r_element = docx.makeelement('r')
				if properties:
					r_element.append(properties.as_element())
				t_element = docx.makeelement('t', tagtext=text)
				if text.strip() != text:
					t_element.attrib["{http://www.w3.org/XML/1998/namespace}space"] = "preserve"
				r_element.append(t_element)
				paragraph_element.append(r_element)

		return [paragraph_element, ]

	def add_hyperlink(self, text, uri):
		self.hyperlink_text = text
		self.hyperlink_uri = uri
		self.hyperlinkId = self.document.add_relationship('http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', self.hyperlink_uri, 'External')
		self.document.has_hyperlink = True


class List(DocumentElement):
	LIST_LETTER_STYLE = "ListUpperLetter"
	DEFAULT_STYLE = LIST_LETTER_STYLE

	def __init__(self, list_items, style=DEFAULT_STYLE, as_html=False, *args, **kwargs):
		self.items = list_items
		self.style = style
		self.num_id = -1
		self.as_html = as_html
		super(List, self).__init__(*args, **kwargs)

	def _parse_formatting(self, html_element, paragraph, formatting={}):
		formatting_tags = {'i': 'italic', 'b': 'bold', 'sup': 'superscript', 'sub': 'subscript'}
		formatting = formatting.copy()
		if isinstance(html_element, bs4.NavigableString):
			paragraph.add_text(html_element.string, **formatting)
			return [(html_element.string, {}),]
		elif isinstance(html_element, bs4.Tag):
			if html_element.name in formatting_tags:
				formatting[formatting_tags[html_element.name]] = True
				for subelement in html_element.contents:
					self._parse_formatting(subelement, paragraph, formatting=formatting)


	def add_styles(self, style_tree, numbering_tree, abstractNum=None):
		# Add a num element to the numbering.xml file.
		num_element = Num(abstractNum=abstractNum, document=self.document)
		num_element.as_element(numbering_tree)

		self.num_id = num_element.id
		self.style = "%s%s" % (self.style, num_element.id)

		# Add a style element to the style.xml file
		style_element = etree.SubElement(style_tree, self.namespace_string("style"))
		style_element.attrib[self.namespace_string("styleId")] = self.style
		style_element.attrib[self.namespace_string("type")] = "paragraph"
		pPr_element = etree.SubElement(style_element, self.namespace_string("pPr"))
		numPr_element = etree.SubElement(pPr_element, self.namespace_string("numPr"))
		numId_element = etree.SubElement(numPr_element, self.namespace_string("numId"))
		numId_element.attrib[self.namespace_string('val')] = str(num_element.id)
		etree.SubElement(pPr_element, self.namespace_string("contextualSpacing"))

			# if element.name == "p":
			# 	p = Paragraph(document=self)
			# 	for subelement in element.contents:
			# 		self._parse_formatting(subelement, p)
			# 	self.add_element(p)
			# else:
			# 	continue

	def get_elements(self):
		to_return = []
		if self.as_html:
			for item in self.items:
				soup = bs4.BeautifulSoup(item)
				for element in soup.body.contents:
					if element.name == "p":
						p = Paragraph(document=self, style=self.style, num_level=0, num_id=self.num_id)
						for subelement in element.contents:
							self.document._parse_formatting(subelement, p)
						to_return += p.get_elements()
		else:
			for item in self.items:
				to_return += Paragraph(text=item, style=self.style, num_level=0, num_id=self.num_id).get_elements()
		return to_return


class Table(DocumentElement):
	DEFAULT_BORDERS = {
		'all': {
			'color': '#000000',
			'sz': 1,
			'space': 0,
			'val': 'single',
		}
	}

	def __init__(self, data, has_heading_row=False, borders=DEFAULT_BORDERS, *args, **kwargs):
		self.data = data
		self.has_heading_row = has_heading_row
		self.borders = borders

		super(Table, self).__init__(*args, **kwargs)

	def get_elements(self):
		return [docx.table(self.data, heading=self.has_heading_row, borders=self.borders), ]


class PageBreak(DocumentElement):
	DEFAULT_TYPE = 'page'
	DEFAULT_ORIENTATION = 'portrait'

	def __init__(self, break_type=DEFAULT_TYPE, orientation=DEFAULT_ORIENTATION,  *args, **kwargs):
		self.type = break_type
		self.orientation = orientation

		super(PageBreak, self).__init__(*args, **kwargs)

	def get_elements(self):
		return [docx.pagebreak(type=self.type, orient=self.orientation), ]


class WordDocument(object):
	PARAGRAPH_STYLE = "BodyText"
	namespace = docx.nsprefixes['w']

	def __init__(self):
		self._document_elements = []
		self._relationships = []
		self._nums = []
		self._abstract_nums = []
		self._document = None
		self.has_hyperlink = False

	def add_element(self, element):
		self._document_elements.append(element)

	def last_element(self):
		try:
			return self._document_elements[-1]
		except IndexError:
			return None

	def add_relationship(self, relationshipNamespace, relationshipTarget, relationshipTargetMode=None):
		relationshipId = len(self._relationships) + 1
		self._relationships.append((relationshipId, [relationshipNamespace, relationshipTarget, relationshipTargetMode]))

		return relationshipId

	def add_heading(self, text, level=Heading.DEFAULT_LEVEL):
		self.add_element(Heading(text, level=level, document=self))
		return self.last_element()

	def insert_pagebreak(self, break_type='page', orientation='portrait'):
		self.add_element(PageBreak(break_type=break_type, orientation=orientation, document=self))
		return self.last_element()

	def add_table(self, table_data, has_heading_row=False, borders=Table.DEFAULT_BORDERS):
		self.add_element(Table(table_data, has_heading_row=has_heading_row, borders=borders, document=self))
		return self.last_element()

	def add_list(self, list_items, style=List.DEFAULT_STYLE, as_html=False):
		self.add_element(List(list_items, style=style, document=self, as_html=as_html))
		return self.last_element()

	def add_list_html(self, list_items, style=List.DEFAULT_STYLE):
		return self.add_list(list_items, style=style, as_html=True)

	def add_paragraph(self, text, style=Paragraph.DEFAULT_STYLE):
		self.add_element(Paragraph(text=text, style=style, document=self))
		return self.last_element()

	def _parse_formatting(self, html_element, paragraph, formatting={}):
		formatting_tags = {'i': 'italic', 'b': 'bold', 'sup': 'superscript', 'sub': 'subscript'}
		formatting = formatting.copy()
		if isinstance(html_element, bs4.NavigableString):
			paragraph.add_text(html_element.string, **formatting)
			return [(html_element.string, {}),]
		elif isinstance(html_element, bs4.Tag):
			if html_element.name in formatting_tags:
				formatting[formatting_tags[html_element.name]] = True
				for subelement in html_element.contents:
					self._parse_formatting(subelement, paragraph, formatting=formatting)

	def add_html(self, html):
		soup = bs4.BeautifulSoup(html)
		for element in soup.body.contents:
			if element.name == "p":
				p = Paragraph(document=self)
				for subelement in element.contents:
					self._parse_formatting(subelement, p)
				self.add_element(p)
			else:
				continue

		return self.last_element()

	def namespace_string(self, tag):
		return "{%s}%s" % (self.namespace, tag)

	def build_styles(self):
		style_template = etree.parse(os.path.join(CURRENT_PATH, 'stylesBase.xml'))
		numbering_template = etree.parse(os.path.join(CURRENT_PATH, 'numberingBase.xml'))
		style_root = style_template.getroot()
		numbering_root = numbering_template.getroot()		

		if self.has_hyperlink:
			style_element = etree.SubElement(style_root,self.namespace_string("style"))
			style_element.attrib[self.namespace_string("styleId")] = "Hyperlink"
			style_element.attrib[self.namespace_string('type')] = "character"
			name_element = etree.SubElement(style_element, self.namespace_string("name"))
			name_element.attrib[self.namespace_string("val")] = "Hyperlink"
			baseOn_element = etree.SubElement(style_element, self.namespace_string("basedOn"))
			baseOn_element.attrib[self.namespace_string('val')] = "DefaultParagraphFont"
			rPr_element = etree.SubElement(style_element, self.namespace_string("rPr"))
			color_element = etree.SubElement(rPr_element, self.namespace_string("color"))
			color_element.attrib[self.namespace_string('val')] = "0000FF"
			color_element.attrib[self.namespace_string('themeColor')] = "hyperlink"
			u_element = etree.SubElement(rPr_element, self.namespace_string("u"))
			u_element.attrib[self.namespace_string('val')] = "single"

		a = AbstractNum(document=self)
		a.add_level(0, start_from=1)
		a.as_element(numbering_root)

		for element in self._document_elements:
			element.add_styles(style_root, numbering_root, abstractNum=a)

		style_outfile = open(os.path.join(docx.template_dir, 'word/styles.xml'), 'w')
		style_template.write(style_outfile, pretty_print=True)
		style_outfile.close()

		numbering_outfile = open(os.path.join(docx.template_dir, 'word/numbering.xml'), 'w')
		numbering_template.write(numbering_outfile, pretty_print=True)
		numbering_outfile.close()

	def build_document(self):
		document = docx.newdocument()
		self.build_styles()
		body = document.xpath('/w:document/w:body', namespaces=docx.nsprefixes)[0]
		for part in self._document_elements:
			for element in part.get_elements():
				body.append(element)

		self._document = document

	def save(self, out):
		self.build_document()
		file_buffer = StringIO.StringIO()
		document_file = zipfile.ZipFile(file_buffer, mode='w', compression=zipfile.ZIP_DEFLATED)

		title = 'Questions'
		subject = 'A set of peer-reviewed MCQ questions for this block.'
		creator = 'Michael Hagarty'
		coreprops = docx.coreproperties(title=title, subject=subject, creator=creator, keywords=[])
		appprops = docx.appproperties()
		contenttypes = docx.contenttypes()
		websettings = docx.websettings()
		relationships = docx.relationshiplist()
		relationships = list(enumerate(relationships, start=len(self._relationships)+1))
		relationships += self._relationships
		#relationships += self._hyperlinks
		wordrelationships = generate_wordrelationships(relationships)

		treesandfiles = {
			self._document: 'word/document.xml',
			coreprops: 'docProps/core.xml',
			appprops: 'docProps/app.xml',
			contenttypes: '[Content_Types].xml',
			websettings: 'word/webSettings.xml',
			wordrelationships: 'word/_rels/document.xml.rels'
		}

		for tree in treesandfiles:
			treestring = etree.tostring(tree, encoding="UTF-8", pretty_print=True, standalone=True)
			document_file.writestr(treesandfiles[tree], treestring)

		files_to_ignore = ['.DS_Store', 'stylesBase.xml', 'numberingBase.xml']
		for dirpath, dirnames, filenames in os.walk(docx.template_dir):
			for filename in filenames:
				if filename in files_to_ignore:
					continue
				templatefile = os.path.join(dirpath, filename)
				archivename = templatefile.replace(docx.template_dir, '')[1:]
				document_file.write(templatefile, archivename)

		document_file.close()

		out.write(file_buffer.getvalue())
		file_buffer.close()
